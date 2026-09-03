import subprocess, sys

def install(pkg):
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg, '--user', '--quiet'])

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    install('python-docx')
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

BLUE = RGBColor(33, 56, 107)
GREY = RGBColor(100, 100, 100)

def section_heading(text):
    h = doc.add_heading(text, level=2)
    run = h.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    return h

def bullet(bold_part, normal_part):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(bold_part + ' ')
    r1.bold = True
    p.add_run(normal_part)

# ------ HEADER ------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Stefan Antonio Cirisanu')
r.bold = True
r.font.size = Pt(22)
r.font.name = 'Arial'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('Software Engineer | Full-Stack & Web3 Architect')
r2.bold = True
r2.font.size = Pt(11)
r2.font.name = 'Arial'

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = contact.add_run('Timisoara, Romania  |  humanityledger@icloud.com  |  linkedin.com/in/stefan-antonio-cirisanu')
r3.font.size = Pt(9)
r3.font.color.rgb = GREY
r3.font.name = 'Arial'

doc.add_paragraph()

# ------ PROFESSIONAL SUMMARY ------
section_heading('Professional Summary')

doc.add_paragraph('Software Engineer specializing in the design, development, and deployment of highly secure, scalable decentralized applications and enterprise web architectures. Recognized for extreme technical autonomy, high-velocity engineering, and the capacity to single-handedly drive massive-scale systems from conceptual architecture to production-ready code.')
doc.add_paragraph('Possess deep expertise across the entire modern TypeScript ecosystem (React, Next.js, Node.js) and EVM-compatible networks (Solidity, Viem, Zero-Knowledge proofs). Adept at managing highly complex, monolithic codebases, implementing zero-trust cryptographic security pipelines, and optimizing global state management for real-time institutional dashboards.')
doc.add_paragraph('Currently relocating to Timisoara, Romania, where he will be delivering university-level lectures in Cryptography and Zero-Knowledge Proofs. Actively seeking demanding IT and Software Engineering roles that require rigorous security standards, robust architectural planning, and end-to-end system ownership.')

# ------ TECHNICAL SKILLS ------
section_heading('Technical Skills')
bullet('Languages:', 'TypeScript, JavaScript (ES6+), Solidity, Java (SE 21), SQL, Python')
bullet('Frontend:', 'React, Next.js (App Router), Tailwind CSS, Framer Motion, Zustand, React Query')
bullet('Backend & DB:', 'Node.js, RESTful APIs, Prisma ORM, PostgreSQL, Redis')
bullet('Web3 & Blockchain:', 'EVM Networks (Ethereum, Base, Optimism, Polygon), Wagmi, Viem, Ethers.js, Hardhat, WalletConnect / AppKit')
bullet('Security & Architecture:', 'Smart Contract Security (Reentrancy, Time-locks), Zero-Knowledge (ZK) Proof integration, EdDSA/HMAC Cryptography, SIWE, JWT lifecycle management.')

# ------ PROFESSIONAL EXPERIENCE ------
section_heading('Professional Experience')

p = doc.add_paragraph()
r = p.add_run('Sole Architect & Lead Engineer')
r.bold = True
r.font.size = Pt(12)

p2 = doc.add_paragraph()
r2 = p2.add_run('Humanity Ledger  |  Remote  |  2023 - Present')
r2.italic = True
r2.font.color.rgb = GREY

bullet('End-to-End System Ownership:', 'Directed the entire software development lifecycle for an institutional-grade Web3 platform. Independently engineered and maintained the frontend, backend, and smart contract layers, sustaining a massive output of high-quality, production-ready code over a single year.')
bullet('Security & Authentication:', 'Designed a zero-trust authentication pipeline utilizing SIWE (Sign-In with Ethereum) alongside secure, HttpOnly cookie-based JWT sessions. Built server-side cryptographic validation for Zero-Knowledge (ZK) proofs to proactively prevent replay attacks and nullifier forgery.')
bullet('Smart Contract Engineering:', 'Developed immutable Solidity smart contracts featuring advanced security mechanics, such as the AegisCircuitBreaker (time-locked withdrawal limits), to protect Total Value Locked (TVL) against rapid algorithmic exploitation.')
bullet('Complex State Management:', 'Built a global settings persistence engine using Zustand, enabling real-time cloud synchronization of intricate UI configurations (dynamic fiat conversion, spam token filtering, encrypted portfolio masking) without triggering unnecessary DOM re-renders.')
bullet('High-Performance Integrations:', 'Integrated complex multi-chain data pipelines using Wagmi and Viem, processing live wallet balances, transaction histories, and real-time ECB exchange rates to power a seamless, low-latency user dashboard.')

# ------ ACADEMIC & TEACHING ------
doc.add_paragraph()
section_heading('Academic & Teaching')

p = doc.add_paragraph()
r = p.add_run('University Lecturer -- Cryptography & Zero-Knowledge Proofs')
r.bold = True
r.font.size = Pt(12)

p2 = doc.add_paragraph()
r2 = p2.add_run('University of Timisoara, Romania  |  Starting December 2026')
r2.italic = True
r2.font.color.rgb = GREY

bullet('', 'Invited to deliver advanced university-level lectures covering cryptographic primitives, applied Zero-Knowledge proof systems, and their real-world implementation in decentralized networks.')
bullet('', 'Subject matter expertise drawn from direct hands-on engineering of ZK-proof verification pipelines in production-grade blockchain applications.')

# ------ EDUCATION ------
doc.add_paragraph()
section_heading('Education & Certifications')

p = doc.add_paragraph()
r = p.add_run('B.Sc. in Cybersecurity Engineering')
r.bold = True

p2 = doc.add_paragraph()
r2 = p2.add_run('Universitat Politecnica de Valencia (UPV) -- Valencia, Spain')
r2.italic = True
r2.font.color.rgb = GREY

bullet('', 'Rigorous academic focus on network security, cryptographic protocols, secure software architecture, and threat mitigation.')

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Java SE 21 Developer')
r.bold = True

p2 = doc.add_paragraph()
r2 = p2.add_run('Professional Training & Certification')
r2.italic = True
r2.font.color.rgb = GREY

bullet('', 'Advanced proficiency in modern Java architecture, concurrent programming, and enterprise-grade backend development.')

# ------ LANGUAGES ------
doc.add_paragraph()
section_heading('Languages')
bullet('Spanish:', 'Native / Bilingual')
bullet('Romanian:', 'Native / Bilingual')
bullet('Valencian:', 'Native / Bilingual')
bullet('English:', 'Professional Working Proficiency')
bullet('Lithuanian:', 'Conversational (Speaking & Listening)')
bullet('Italian:', 'Basic Comprehension (Listening & Reading)')
bullet('German:', 'Elementary / Basic')

out_path = 'd:/Projects/Wallet Human Polymarket ID/CV_Stefan_Antonio_Cirisanu.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
